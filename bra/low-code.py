{
 "cells": [
  {
   "cell_type": "code",
   "execution_count": 1,
   "metadata": {},
   "outputs": [
    {
     "ename": "SyntaxError",
     "evalue": "incomplete input (1397751096.py, line 26)",
     "output_type": "error",
     "traceback": [
      "\u001b[1;36m  Cell \u001b[1;32mIn[1], line 26\u001b[1;36m\u001b[0m\n\u001b[1;33m    \u001b[0m\n\u001b[1;37m    ^\u001b[0m\n\u001b[1;31mSyntaxError\u001b[0m\u001b[1;31m:\u001b[0m incomplete input\n"
     ]
    }
   ],
   "source": [
    "from docx import Document\n",
    "\n",
    "# Criação do documento Word\n",
    "doc = Document()\n",
    "\n",
    "# Título do documento\n",
    "doc.add_heading('Governança de Low-Code com Power Apps', level=1)\n",
    "\n",
    "# Adicionando os conteúdos em texto estruturado\n",
    "doc.add_paragraph(\n",
    "    \"A governança eficiente de low-code com Power Apps é essencial para garantir segurança, conformidade e eficiência \"\n",
    "    \"no uso da plataforma. O primeiro passo é estabelecer uma estrutura de governança que defina claramente os papéis e \"\n",
    "    \"responsabilidades das diferentes partes envolvidas, incluindo os desenvolvedores cidadãos, a equipe de TI e os \"\n",
    "    \"desenvolvedores profissionais. A criação de uma equipe de governança composta por membros da TI, compliance e \"\n",
    "    \"áreas de negócios também é fundamental para supervisionar o uso adequado da plataforma.\"\n",
    ")\n",
    "\n",
    "doc.add_paragraph(\n",
    "    \"É importante desenvolver políticas de uso que limitem permissões e controlem quem pode criar e publicar aplicativos. \"\n",
    "    \"Além disso, deve-se controlar o uso de conectores, principalmente em relação a dados sensíveis, e estabelecer regras \"\n",
    "    \"de nomenclatura para facilitar a organização e o gerenciamento dos aplicativos criados.\"\n",
    ")\n",
    "\n",
    "doc.add_paragraph(\n",
    "    \"A implementação de ferramentas de monitoramento e auditoria é uma etapa crucial. Utilizar o Power Platform Admin \"\n",
    "   \n"
   ]
  },
  {
   "cell_type": "code",
   "execution_count": null,
   "metadata": {},
   "outputs": [],
   "source": []
  }
 ],
 "metadata": {
  "kernelspec": {
   "display_name": "Python 3",
   "language": "python",
   "name": "python3"
  },
  "language_info": {
   "codemirror_mode": {
    "name": "ipython",
    "version": 3
   },
   "file_extension": ".py",
   "mimetype": "text/x-python",
   "name": "python",
   "nbconvert_exporter": "python",
   "pygments_lexer": "ipython3",
   "version": "3.10.11"
  }
 },
 "nbformat": 4,
 "nbformat_minor": 2
}
